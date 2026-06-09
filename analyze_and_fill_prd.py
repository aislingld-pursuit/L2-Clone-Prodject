"""Fill PRD template in-place preserving original Word layout and styles."""
from __future__ import annotations

import copy
import shutil
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph

BASE = Path(r"C:\Users\Aisling Ld Pursuit\OneDrive\Documents\Pursuit L2 Project\L2 Clone of Wisper")
SRC = BASE / "Aisling Copy of 20260515 PRD Template.docx"
WORK = BASE / "_prd_working_copy.docx"
OUT = BASE / "Jimmy and Aisling Copy of 20260515 PRD Template - FILLED.docx"

DIRECTION_FILL = "D6E4F0"
HEADER_FILL = "1F4E8C"

REPLACEMENTS = {
    "[BUILD NAME]": "WISPER — LOCAL WHISPER NOTES CLONE",
    "[Build name]": "Wisper (Local-First Whisper Notes Clone)",
    "[Your name(s)]": "Aisling Ld Pursuit",
    "[Date]": "June 8, 2026",
    "[Write your problem description here.]": (
        "Knowledge workers, students, journalists, and meeting participants experience "
        "lost information and slow capture of spoken content in meetings, lectures, "
        "interviews, and voice notes because speech is ephemeral and manual typing is "
        "3–4× slower than speaking, resulting in missed details, poor searchability, "
        "and hours spent re-listening to recordings or reconstructing notes from memory."
    ),
    "[Insert data point or research finding]": (
        "Most people speak 120–160 words per minute but type only ~40 wpm, creating a "
        "persistent capture bottleneck (industry benchmarks on dictation vs typing speed)."
    ),
    "[Insert user pain point observed or validated]": (
        "Users report losing action items, numbers, and names during live meetings because "
        "they cannot simultaneously participate and take complete notes."
    ),
    "[Insert market insight or competitive context]": (
        "Whisper Notes proves demand for a privacy-first, offline transcription app using Whisper "
        "Large V3 Turbo on-device—but it is Apple-only. Most competitors default to cloud APIs, "
        "creating a gap for a cross-platform, local-first Whisper Notes clone."
    ),
    "[Write your opportunity statement here.]": (
        "Build a full local clone of Whisper Notes: record or import audio, transcribe entirely "
        "on-device with Whisper-class accuracy, organize in a searchable library, and export "
        "text/subtitles—without accounts, cloud upload, or network permissions for core features."
    ),
    "[Insert market size or growth data]": (
        "Global speech-to-text market projected to exceed $5B by 2030, driven by remote work, "
        "accessibility requirements, and AI adoption."
    ),
    "[Insert link to strategy doc or competitive analysis, if available]": (
        "See Appendix for competitive reference apps, technical constraints, and open questions. "
        "Competitive feature matrix and Figma wireframes to be added in follow-up artifacts."
    ),
    "[Who they are and what they care about]": (
        "Primary users are professionals and students who regularly capture spoken content—"
        "meetings, lectures, interviews, podcasts, and voice memos—and need accurate transcripts "
        "they can search, edit, and reuse quickly."
    ),
    "[Who else interacts with this product, if anyone]": (
        "Secondary users include content creators, journalists, researchers, and "
        "accessibility-focused users who need captions, subtitles, or translated text from "
        "audio sources."
    ),
    "[Write your solution description here.]": (
        "Wisper is a privacy-first, local-only speech-to-text application (Whisper Notes clone) "
        "for Windows, macOS, Linux, iOS, and Android. Users record voice, import local audio/video, "
        "or paste a YouTube (or other media) URL—audio is downloaded to the device, then transcribed "
        "entirely on-device via whisper.cpp (Whisper Large V3 Turbo) with no cloud STT upload. "
        "Results live in a local SQLite library with full-text search, inline editing, and export "
        "to TXT, SRT, and VTT."
    ),
    "[Write your value proposition here.]": (
        "Privacy-conscious professionals, students, journalists, and clinicians who cannot send "
        "sensitive audio to the cloud use Wisper to record, transcribe, search, and export "
        "spoken content entirely on their own hardware. Unlike cloud-first transcription apps "
        "that require upload and accounts, Wisper matches Whisper Notes’ offline accuracy and "
        "simplicity while supporting every major desktop and mobile platform."
    ),
    "[Write your vitamin value prop]": (
        "Reliable audio capture and file import with accurate transcription into readable, "
        "editable text—the baseline users expect from any serious transcription app."
    ),
    "[Write your painkiller value prop]": (
        "Eliminates the need to manually type or re-listen to recordings by producing searchable "
        "transcripts automatically from any recording or import."
    ),
    "[Your steroid value prop]": (
        "Paste a YouTube lecture link or drop a folder of recordings—audio is pulled or imported "
        "locally, transcribed entirely on-device, and searchable in your library before you finish "
        "your coffee. No cloud transcription, ever."
    ),
    "[What you’re trying to achieve — frame as an outcome, not a feature]": (
        "Give users a single app to capture, transcribe, search, and export spoken content "
        "without switching tools."
    ),
    "[Business or user benefit you’re optimizing for]": (
        "Deliver on-device transcription accuracy matching Whisper Notes (Whisper Large V3 Turbo) "
        "with zero outbound network traffic during core transcription workflows."
    ),
    "[Success outcome for this version]": (
        "Ship a desktop MVP where users complete record, local import, or YouTube URL → download → "
        "local transcribe → search → export—with zero cloud STT and a local transcript library "
        "as the core retention surface."
    ),
    "[What’s explicitly out of scope for this version, and why]": (
        "Cloud transcription APIs, user accounts, and cloud sync—excluded because transcription "
        "must stay on-device. (YouTube/URL import is in scope: network is used only to download "
        "media; transcription remains local.)"
    ),
    "[What you’re NOT building right now]": (
        "Cloud sync, team workspaces, enterprise SSO, real-time live captions during recording, "
        "and speaker diarization—deferred until single-user local desktop MVP is complete."
    ),
    "[Goal from section 2b]": "Core transcription value",
    "[Measurable indicator]": "Users complete a transcription end-to-end",
    "[How you’ll track it]": "Transcription completion rate",
    "[Specific number or threshold]": ">75% of started jobs",
    "[Add supporting links, research, wireframes, or open questions here.]": (
        "__APPENDIX_PLACEHOLDER__"
    ),
}

NEEDS = [
    "As a privacy-conscious user, I need all transcription to happen on my device because sensitive audio must never be uploaded to a cloud service.",
    "As a meeting participant, I need to record and transcribe conversations locally because I cannot take notes while actively contributing and cannot risk cloud exposure.",
    "As a student, I need to paste a YouTube lecture URL and get a local transcript because most course content lives on YouTube and I will not upload it to a cloud STT service.",
    "As a cross-platform user, I need the same local-first experience on Windows, Mac, and Linux because Whisper Notes is Apple-only today.",
    "As a content creator, I need to export SRT/VTT subtitles from local transcripts because I edit video without sending audio to third parties.",
]

GOALS_EXTRA = [
    "Achieve Whisper Notes–class accuracy using Whisper Large V3 Turbo running locally via whisper.cpp.",
    "Ship YouTube URL import (yt-dlp download → local transcribe) as a P0 MVP feature alongside record and file import.",
    "Support Windows, macOS, and Linux desktop as MVP; iOS and Android in subsequent phases.",
]

NON_GOALS_EXTRA = [
    "Cloud transcription, cloud sync, user accounts, and any remote STT API—never send audio or transcripts to third-party transcription services.",
]

METRICS = [
    ("Core local transcription", "Users complete transcription without cloud STT", "Transcription completion rate", ">90% of started jobs"),
    ("YouTube import (P0)", "User pastes URL → local transcript", "YouTube URL import success rate", ">85% of valid public URLs"),
    ("Privacy guarantee", "No audio sent to STT cloud during transcribe", "Outbound STT API calls during transcription", "0"),
    ("Library value", "Users find past transcripts", "Search latency across library", "<100ms for 1,000 items"),
]

JOURNEYS = [
    {
        "title": "User Journey 1: Professional capturing a live meeting or dictation",
        "context": (
            "This is the highest-frequency path. Users need low-friction recording and fast access "
            "to transcript text during or immediately after a conversation."
        ),
        "subs": [
            ("Sub-journey: Starting and recording audio", [
                ("P0", "User can start a new recording from the home screen without creating an account."),
                ("P0", "Home screen exposes three equal entry points: Record, Import file, Paste YouTube URL."),
                ("P0", "User can pause and resume an in-progress recording."),
                ("P0", "User can see recording duration and input level while recording."),
                ("P1", "User can add a title or tags to a recording before or after capture."),
                ("P2", "User can choose audio quality settings (standard vs high fidelity)."),
            ]),
            ("Sub-journey: Live and post-recording transcription", [
                ("P0", "User can transcribe a completed recording on-device with timestamps (batch processing)."),
                ("P0", "User can select transcription language or use auto-detect (99+ languages)."),
                ("P0", "Transcription runs with zero outbound network connections."),
                ("P1", "User can view progress (percent complete, ETA) while transcription runs locally."),
                ("P1", "User can retry transcription if processing fails."),
                ("P2", "User can enable live partial transcription during recording (deferred; Whisper Notes is batch-first)."),
            ]),
            ("Sub-journey: Reviewing and editing transcript", [
                ("P0", "User can view transcript text synced to timestamps."),
                ("P0", "User can edit transcript text inline and save changes."),
                ("P0", "User can search within a transcript for keywords."),
                ("P1", "User can jump to a timestamp by clicking a transcript segment."),
                ("P2", "User can add manual bookmarks/notes at specific timestamps."),
            ]),
        ],
    },
    {
        "title": "User Journey 2: Student or researcher importing existing audio/video",
        "context": (
            "Many users already have lecture recordings, interview files, podcast audio, or YouTube "
            "links. Import must support local files and URL-based fetch, with transcription always "
            "on-device after audio is on the machine."
        ),
        "subs": [
            ("Sub-journey: Importing media", [
                ("P0", "User can import audio files (MP3, M4A, WAV, FLAC) from local storage."),
                ("P0", "User can import video files (MP4, MOV) and extract audio locally for transcription."),
                ("P0", "User can paste a YouTube URL on the home screen and download audio locally for transcription (yt-dlp)."),
                ("P0", "User can see download progress, then local transcription progress, as separate steps."),
                ("P1", "User can paste other supported media URLs that yt-dlp handles (e.g. direct audio/video links, podcasts)."),
                ("P1", "User can drag-and-drop multiple local files into the app."),
                ("P2", "User can batch-import multiple files or URLs into a sequential local queue."),
            ]),
            ("Sub-journey: Processing imported media", [
                ("P0", "User can transcribe imported files using the same language and quality options as live recordings."),
                ("P0", "User can view processing status (queued, processing, complete, failed)."),
                ("P1", "User can cancel an in-progress transcription job."),
                ("P2", "User can auto-split long files into chapter segments."),
            ]),
        ],
    },
    {
        "title": "User Journey 3: User organizing, searching, and reusing transcripts",
        "context": (
            "The transcript library is the retention engine. Users must find past content quickly and "
            "export it into existing workflows."
        ),
        "subs": [
            ("Sub-journey: Library and search", [
                ("P0", "User can view a list of all recordings/transcripts sorted by date."),
                ("P0", "User can search across all transcripts by keyword and see matching results."),
                ("P0", "User can open any saved transcript from the library."),
                ("P1", "User can filter library items by date, language, or tag."),
                ("P2", "User can archive or delete transcripts."),
            ]),
            ("Sub-journey: Export and sharing", [
                ("P0", "User can export transcript as plain text (.txt)."),
                ("P0", "User can copy full transcript or selected text to clipboard."),
                ("P0", "User can export subtitles as SRT or VTT with timestamps."),
                ("P1", "User can export transcript as Markdown with metadata."),
                ("P1", "User can export to a user-chosen local folder (no cloud required)."),
                ("P2", "User can export DOCX or PDF from local transcript."),
            ]),
        ],
    },
    {
        "title": "User Journey 4: User getting local AI insights from long recordings (v2)",
        "context": (
            "Optional local LLM summaries (llama.cpp) may be added after core transcription ships. "
            "All processing remains on-device—no cloud LLM APIs."
        ),
        "subs": [
            ("Sub-journey: Local summaries and highlights", [
                ("P2", "User can generate an on-device summary of a completed transcript (local LLM)."),
                ("P2", "User can view bullet-point highlights of key moments."),
                ("P2", "User can regenerate summary with different length (short vs detailed)."),
            ]),
            ("Sub-journey: Chat with transcript (local only)", [
                ("P2", "User can ask questions about a transcript answered by a local model."),
                ("P2", "User can see answers with referenced timestamps from local transcript."),
            ]),
        ],
    },
    {
        "title": "User Journey 5: Privacy-first architecture verification",
        "context": (
            "Privacy is enforced by architecture: transcription never leaves the device. Network is "
            "used only when the user explicitly imports from a URL (e.g. YouTube) or downloads the "
            "Whisper model. No accounts or cloud STT services."
        ),
        "subs": [
            ("Sub-journey: Local-only transcription", [
                ("P0", "All transcription runs on-device via whisper.cpp (Whisper Large V3 Turbo)."),
                ("P0", "Audio and transcripts are never sent to a cloud transcription API."),
                ("P0", "User is never required to create an account or sign in."),
                ("P0", "Record, local-file import, and YouTube URL import work after the Whisper model is installed."),
                ("P0", "User can see whether an item used network (URL download) vs was fully offline (record/local file)."),
                ("P1", "User can delete original audio after transcription while keeping text locally."),
                ("P1", "User can choose model size (quality vs speed) for their hardware."),
                ("P2", "User can verify in settings that no analytics SDK sends audio or transcript data."),
            ]),
        ],
    },
    {
        "title": "User Journey 6: First-run setup and settings (no cloud account)",
        "context": (
            "First-run focuses on downloading/bundling the local model and granting microphone "
            "permission. No signup, sync, or subscription backend."
        ),
        "subs": [
            ("Sub-journey: First-run experience", [
                ("P0", "User can complete first local transcription without signing up."),
                ("P0", "User can download or install the Whisper model on first launch (one-time)."),
                ("P0", "User can grant microphone permission and see clear error states if denied."),
                ("P1", "User can complete short onboarding: model install → record → transcribe → export."),
                ("P2", "User can skip onboarding and access app immediately."),
            ]),
            ("Sub-journey: Settings and data control", [
                ("P1", "User can set default transcription language and model quality."),
                ("P1", "User can choose app data directory and export folder."),
                ("P2", "User can export full library backup to a local `.wisper` bundle."),
            ]),
        ],
    },
]


def paragraph_fill(p: Paragraph) -> str | None:
    pPr = p._element.pPr
    if pPr is None:
        return None
    shd = pPr.find(qn("w:shd"))
    if shd is None:
        return None
    return (shd.get(qn("w:fill")) or "").upper()


def set_paragraph_text(p: Paragraph, text: str) -> None:
    if p.runs:
        p.runs[0].text = text
        for r in p.runs[1:]:
            r.text = ""
    else:
        p.text = text


def replace_placeholders(p: Paragraph) -> None:
    text = p.text
    if not text:
        return
    new = text
    for old, val in REPLACEMENTS.items():
        if old in new:
            new = new.replace(old, val)
    if new != text:
        set_paragraph_text(p, new)


def delete_paragraph(p: Paragraph) -> None:
    el = p._element
    el.getparent().remove(el)


def clone_after(ref: Paragraph, text: str, style=None) -> Paragraph:
    new_el = copy.deepcopy(ref._element)
    ref._element.addnext(new_el)
    new_p = Paragraph(new_el, ref._parent)
    set_paragraph_text(new_p, text)
    if style is not None:
        new_p.style = style
    return new_p


APPENDIX_LINES = [
    "Competitive Reference Apps",
    "• Whisper Notes (iOS/macOS) — PRIMARY CLONE TARGET: offline Whisper Large V3 Turbo, Core ML, local library, SRT/VTT export, no subscriptions, no network for transcription.",
    "• Whisper Transcription (iOS): local + cloud options — we match local path only, cross-platform.",
    "• Whisper – Speech to Text (iOS): cloud-heavy — explicitly NOT our model.",
    "",
    "Technical Constraints (Local-First MVP)",
    "• Transcription engine: whisper.cpp with Whisper Large V3 Turbo (GGML/GGUF); WhisperKit on Apple where beneficial.",
    "• Target platforms: Windows, macOS, Linux (Tauri 2 desktop MVP); iOS and Android in Phase 5–6.",
    "• Privacy: zero outbound network during transcription; SQLite local storage; no user accounts.",
    "• Supported import (P0): local MP3/M4A/WAV/FLAC/MP4/MOV; YouTube URL via yt-dlp (download only—transcription stays local).",
    "• See TECHNICAL_ARCHITECTURE.md and ROADMAP.md in project folder.",
    "",
    "Open Questions",
    "• Model delivery: bundle in installer vs one-time download on first launch?",
    "• Intel Mac / low-RAM Windows: ship smaller default model (small.en) with upgrade path?",
    "• Pricing: one-time purchase (Whisper Notes model) vs free/open-source?",
    "• YouTube: ship yt-dlp bundled vs require user install? Mobile URL import timing (iOS/Android)?",
    "",
    "Suggested Next Artifacts",
    "• TECHNICAL_ARCHITECTURE.md — completed.",
    "• ROADMAP.md — completed (agent-paced estimates).",
    "• User flow wireframes: Record, Import, Library, Transcript Detail, Export.",
    "• Phase 0 spike: Tauri + whisper.cpp end-to-end transcribe.",
]


def fill_appendix(doc: Document) -> None:
    anchor = None
    for p in doc.paragraphs:
        if "__APPENDIX_PLACEHOLDER__" in (p.text or ""):
            anchor = p
            break
    if anchor is None:
        return
    set_paragraph_text(anchor, APPENDIX_LINES[0])
    insert = anchor
    for line in APPENDIX_LINES[1:]:
        insert = clone_after(insert, line, anchor.style)


def fill_user_needs(doc: Document) -> None:
    for i, p in enumerate(doc.paragraphs):
        if (p.text or "").strip() == "Key User Needs":
            idx = i
            break
    if idx is None:
        return

    need_paras: list[Paragraph] = []
    for p in doc.paragraphs[idx + 1 :]:
        if p.style and p.style.name.startswith("Heading"):
            break
        t = (p.text or "").strip()
        if t:
            need_paras.append(p)

    if not need_paras:
        return

    set_paragraph_text(need_paras[0], NEEDS[0])
    insert = need_paras[0]
    for i, need in enumerate(NEEDS[1:], start=1):
        if i < len(need_paras):
            set_paragraph_text(need_paras[i], need)
        else:
            insert = clone_after(insert, need, need_paras[0].style)
    for extra in need_paras[len(NEEDS) :]:
        delete_paragraph(extra)


def fill_goals_and_non_goals(doc: Document) -> None:
    goals = [
        "Give users a single local app to capture, transcribe, search, and export spoken content without cloud STT or accounts.",
        "Deliver Whisper Notes–class on-device accuracy using Whisper Large V3 Turbo via whisper.cpp.",
        "Ship YouTube URL import (yt-dlp → local transcribe) as P0 alongside record and file import.",
        "Ship cross-platform desktop MVP (Windows, macOS, Linux) with local SQLite library and search.",
        "Maintain zero outbound connections to cloud STT APIs during all transcription workflows.",
    ]
    non_goals = [
        "Cloud transcription APIs (OpenAI Whisper API or any remote STT)—excluded by privacy-first architecture.",
        "User accounts, cloud sync, and subscription backends—excluded from v1.",
        "Real-time live captions during recording and speaker diarization—deferred until batch local MVP is complete.",
    ]

    def fill_block(heading: str, lines: list[str]) -> None:
        idx = None
        for i, p in enumerate(doc.paragraphs):
            if (p.text or "").strip() == heading:
                idx = i
                break
        if idx is None:
            return
        paras: list[Paragraph] = []
        for p in doc.paragraphs[idx + 1 :]:
            if p.style and p.style.name.startswith("Heading"):
                break
            if (p.text or "").strip():
                paras.append(p)
        if not paras:
            return
        set_paragraph_text(paras[0], lines[0])
        insert = paras[0]
        for i, line in enumerate(lines[1:], start=1):
            if i < len(paras):
                set_paragraph_text(paras[i], line)
            else:
                insert = clone_after(insert, line, paras[0].style)
        for extra in paras[len(lines) :]:
            delete_paragraph(extra)

    fill_block("Goals", goals)
    fill_block("Non-Goals", non_goals)


def fill_metrics_table(doc: Document) -> None:
    for table in doc.tables:
        hdr = [c.text.strip() for c in table.rows[0].cells]
        if hdr != ["Goal", "Signal", "Metric", "Target"]:
            continue
        while len(table.rows) < len(METRICS) + 1:
            table.add_row()
        for ri, row in enumerate(METRICS, start=1):
            for ci, val in enumerate(row):
                cell = table.rows[ri].cells[ci]
                if cell.paragraphs:
                    set_paragraph_text(cell.paragraphs[0], val)
                    for extra in cell.paragraphs[1:]:
                        set_paragraph_text(extra, "")
                else:
                    cell.text = val


def replace_requirements(doc: Document) -> None:
    req_idx = start_idx = end_idx = None
    for i, p in enumerate(doc.paragraphs):
        t = (p.text or "").strip()
        if t == "3. REQUIREMENTS":
            req_idx = i
        if start_idx is None and t.startswith("User Journey 1:"):
            start_idx = i
        if t.startswith("4. APPENDIX"):
            end_idx = i
            break
    if req_idx is None or end_idx is None:
        return

    # Capture style references from first journey placeholder if present
    if start_idx is not None and start_idx < end_idx:
        journey_style = doc.paragraphs[start_idx].style
        body_ref = doc.paragraphs[start_idx + 1] if start_idx + 1 < end_idx else doc.paragraphs[start_idx]
        body_style = body_ref.style
        for i in range(end_idx - 1, start_idx - 1, -1):
            delete_paragraph(doc.paragraphs[i])
    else:
        # No placeholder block; use paragraph after REQUIREMENTS heading
        journey_style = doc.paragraphs[req_idx].style
        body_ref = doc.paragraphs[req_idx]
        body_style = body_ref.style
        for i in range(end_idx - 1, req_idx, -1):
            delete_paragraph(doc.paragraphs[i])

    appendix = None
    for p in doc.paragraphs:
        if (p.text or "").strip().startswith("4. APPENDIX"):
            appendix = p
            break
    if appendix is None:
        return

    insert_after = doc.paragraphs[req_idx]
    for journey in JOURNEYS:
        insert_after = clone_after(insert_after, journey["title"], journey_style)
        insert_after = clone_after(insert_after, f"Context:  {journey['context']}", body_style)
        insert_after = clone_after(insert_after, "", body_style)
        for sub_title, reqs in journey["subs"]:
            insert_after = clone_after(insert_after, sub_title, body_style)
            for pri, req in reqs:
                insert_after = clone_after(insert_after, f"[{pri}]  {req}", body_style)
            insert_after = clone_after(insert_after, "", body_style)


def remove_direction_boxes(doc: Document) -> None:
    to_delete = []
    for p in doc.paragraphs:
        fill = paragraph_fill(p)
        if fill == DIRECTION_FILL:
            to_delete.append(p)
    for p in reversed(to_delete):
        delete_paragraph(p)


def main() -> None:
    if not SRC.exists() and not WORK.exists():
        raise SystemExit(f"Template not found: {SRC}")
    if SRC.exists():
        try:
            shutil.copy2(SRC, WORK)
        except PermissionError:
            if not WORK.exists():
                raise SystemExit(
                    f"Cannot read template (file may be open): {SRC}. "
                    f"Close it in Word or use existing {WORK}."
                )
            print(f"Warning: using existing working copy (template locked): {WORK}")

    doc = Document(str(WORK))

    for p in doc.paragraphs:
        replace_placeholders(p)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    replace_placeholders(p)

    fill_user_needs(doc)
    fill_goals_and_non_goals(doc)
    fill_appendix(doc)

    fill_metrics_table(doc)
    remove_direction_boxes(doc)
    replace_requirements(doc)

    doc.save(str(OUT))
    print(f"Saved: {OUT}")


if __name__ == "__main__":
    main()
