"""Dump paragraph styles and structure from template copy."""
from pathlib import Path
from docx import Document
from docx.oxml.ns import qn

path = Path(r"C:\Users\Aisling Ld Pursuit\OneDrive\Documents\Pursuit L2 Project\L2 Clone of Wisper\Aisling Copy of 20260515 PRD Template - FILLED.docx")

doc = Document(str(path))
print("=== PARAGRAPH DUMP ===")
for i, p in enumerate(doc.paragraphs):
    t = (p.text or "").replace("\n", " ")[:120]
    style = p.style.name if p.style else "None"
    # check shading
    shd = ""
    pPr = p._element.pPr
    if pPr is not None:
        shd_el = pPr.find("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}shd")
        if shd_el is not None:
            shd = f" fill={shd_el.get(qn('w:fill'))}"
    print(f"{i:3d} | {style:20s}{shd} | {t}")

print("\n=== TABLES ===")
for ti, table in enumerate(doc.tables):
    print(f"Table {ti}: {len(table.rows)} rows x {len(table.columns)} cols")
    for ri, row in enumerate(table.rows[:6]):
        print("  ", " | ".join(c.text[:40].replace('\n',' ') for c in row.cells))
