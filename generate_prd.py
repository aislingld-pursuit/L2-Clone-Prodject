"""Generate filled PRD docx for Wisper clone."""
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_BREAK

OUT = r"C:\Users\Aisling Ld Pursuit\OneDrive\Documents\Pursuit L2 Project\L2 Clone of Wisper\Wisper Clone PRD - Filled.docx"

doc = Document()

def h(text, level=1):
    doc.add_heading(text, level=level)

def p(text):
    doc.add_paragraph(text)

def bullet(text):
    doc.add_paragraph(text, style="List Bullet")

# Title block
h("Product Requirements Document: Net New Build", 0)
p("Build name: Wisper Clone")
p("Owner: Aisling Ld Pursuit")
p("Date: June 8, 2026")

# 1. PROBLEM
h("1. PROBLEM", 1)
p(
    "Knowledge workers, students, journalists, and meeting participants experience "
    "lost information and slow capture of spoken content in meetings, lectures, "
    "interviews, and voice notes because speech is ephemeral and manual typing is "
    "3–4× slower than speaking, resulting in missed details, poor searchability, "
    "and hours spent re-listening to recordings or reconstructing notes from memory."
)

h("Supporting Context", 2)
bullet("Most people speak 120–160 words per minute but type only ~40 wpm, creating a persistent capture bottleneck.")
bullet("Users of existing transcription tools report friction around privacy (cloud uploads), cost (subscriptions), and workflow gaps (import → transcribe → find → export).")
bullet("The speech-to-text app market is growing, but many solutions split features across recording, transcription, search, and AI summarization without a unified, fast experience.")

h("1a. Opportunity", 2)
p(
    "Build a unified speech-to-text product that turns any spoken moment—live dictation, "
    "recordings, or imported audio/video—into accurate, searchable, editable text with "
    "optional AI summaries, capturing demand from professionals who need fast, trustworthy "
    "transcription without a fragmented toolchain."
)
p("Market Opportunity")
bullet("Global speech-to-text market projected to exceed $5B by 2030, driven by remote work, accessibility, and AI adoption.")
bullet("Whisper-class apps on mobile/desktop show strong demand for offline-capable, multilingual transcription with export and search.")

h("1b. Users & Needs", 2)
p("Primary user(s): Professionals and students who regularly capture spoken content—meetings, lectures, interviews, podcasts, and voice memos—and need accurate transcripts they can search, edit, and reuse.")
p("Secondary users: Content creators, journalists, researchers, and accessibility-focused users who need captions, subtitles, or translated text from audio sources.")
p("Key User Needs")
bullet("As a meeting participant, I need to record and transcribe conversations because I cannot reliably take notes while actively contributing.")
bullet("As a student, I need to import lecture recordings and get searchable transcripts because reviewing audio alone is slow and inefficient.")
bullet("As a privacy-conscious user, I need control over where my audio is processed because sensitive conversations should not require cloud upload.")
bullet("As a multilingual user, I need transcription and optional translation across many languages because my work spans more than one language.")
bullet("As a busy professional, I need AI-generated summaries and highlights because I do not have time to read full transcripts of long recordings.")

# 2. PROPOSED SOLUTION
h("2. PROPOSED SOLUTION", 1)
p(
    "Wisper is a cross-platform speech-to-text application that records live audio, "
    "imports audio/video files, and converts speech into accurate, timestamped text "
    "using OpenAI Whisper–class transcription. Users simply press record or import a "
    "file, and the system transcribes, stores, and organizes the result in a searchable "
    "library with editing, export, and optional AI summary tools. As a result, they can "
    "capture spoken information as quickly as they speak it, find any moment by keyword, "
    "and reuse content in documents, captions, or summaries without re-listening."
)

h("2a. Value Proposition", 2)
p(
    "Professionals and students who struggle with slow manual note-taking and hard-to-search "
    "audio recordings use Wisper, a speech-to-text app, to record, transcribe, search, and "
    "summarize spoken content in one place. Unlike generic voice memo apps or fragmented "
    "cloud-only transcription services, it combines fast Whisper-based transcription, a "
    "unified transcript library, and optional on-device processing, helping users turn "
    "speech into usable text in minutes instead of hours."
)

h("2b. Top 3 MVP Value Props", 2)
p("The Vitamin (must-have baseline): Reliable audio capture and file import with accurate transcription into readable, editable text.")
p("The Painkiller (solves the core pain): Eliminates the need to manually type or re-listen to recordings by producing searchable transcripts automatically.")
p("The Steroid (the magic moment): Paste a YouTube link or drop a meeting recording and get a timestamped transcript plus an AI summary of key points in seconds.")

h("2c. Goals & Non-Goals", 2)
p("Goals")
bullet("Give users a single app to capture, transcribe, search, and export spoken content without switching tools.")
bullet("Deliver transcription accuracy comparable to leading Whisper-based apps for English and at least 10 additional MVP languages.")
bullet("Reduce time from recording/import to usable transcript to under 2 minutes for a 30-minute file on target hardware.")
bullet("Achieve 40% week-2 retention among users who complete at least one successful transcription.")
bullet("Establish a searchable transcript library as the core retention surface, not one-off exports.")

p("Non-Goals")
bullet("Full speaker diarization (who said what) for multi-party calls in MVP—deferred to v2 due to model complexity.")
bullet("Real-time collaborative editing or team workspaces—deferred until single-user workflow is validated.")
bullet("Native mobile apps for iOS/Android in MVP—initial release targets web + desktop to accelerate core transcription loop.")
bullet("Custom model training or enterprise SSO/admin console—out of scope for first release.")

h("2d. Success Metrics", 2)
table = doc.add_table(rows=5, cols=4)
table.style = "Table Grid"
headers = ["Goal", "Signal", "Metric", "Target"]
for i, header in enumerate(headers):
    table.rows[0].cells[i].text = header

rows = [
    ("Core transcription value", "Users complete a transcription end-to-end", "Transcription completion rate", ">75% of started jobs"),
    ("Speed to value", "Users get transcript quickly", "Median time from import/record to transcript ready", "<2 min for 30-min file"),
    ("Engagement", "Users return to search or reuse transcripts", "Weekly active users (WAU)", ">500 WAU at 30 days post-launch"),
    ("Retention", "Users find ongoing value in library", "Week-2 retention (completed ≥1 transcription)", ">40%"),
]
for r_idx, row in enumerate(rows, start=1):
    for c_idx, val in enumerate(row):
        table.rows[r_idx].cells[c_idx].text = val

# 3. REQUIREMENTS
h("3. REQUIREMENTS", 1)

# Journey 1
h("User Journey 1: Professional capturing a live meeting or dictation", 2)
p("Context: This is the highest-frequency path. Users need low-friction recording and fast access to transcript text during or immediately after a conversation.")

h("Sub-journey: Starting and recording audio", 3)
bullet("[P0] User can start a new recording from the home screen without creating an account.")
bullet("[P0] User can pause and resume an in-progress recording.")
bullet("[P0] User can see recording duration and input level while recording.")
bullet("[P1] User can add a title or tags to a recording before or after capture.")
bullet("[P2] User can choose audio quality settings (standard vs high fidelity).")

h("Sub-journey: Live and post-recording transcription", 3)
bullet("[P0] User can transcribe a completed recording into text with timestamps.")
bullet("[P0] User can select transcription language or use auto-detect.")
bullet("[P1] User can view partial/streaming transcript output while transcription is in progress.")
bullet("[P1] User can retry transcription if processing fails.")
bullet("[P2] User can enable live caption-style transcription during recording.")

h("Sub-journey: Reviewing and editing transcript", 3)
bullet("[P0] User can view transcript text synced to timestamps.")
bullet("[P0] User can edit transcript text inline and save changes.")
bullet("[P0] User can search within a transcript for keywords.")
bullet("[P1] User can jump to a timestamp by clicking a transcript segment.")
bullet("[P2] User can add manual bookmarks/notes at specific timestamps.")

# Journey 2
h("User Journey 2: Student or researcher importing existing audio/video", 2)
p("Context: Many users already have lecture recordings, interview files, or podcast audio. Import must support common formats and feel as seamless as recording in-app.")

h("Sub-journey: Importing media", 3)
bullet("[P0] User can import audio files (MP3, M4A, WAV) from local storage.")
bullet("[P0] User can import video files (MP4, MOV) and extract audio for transcription.")
bullet("[P1] User can paste a YouTube URL and import audio for transcription.")
bullet("[P1] User can see import progress and estimated processing time.")
bullet("[P2] User can batch-import multiple files into a queue.")

h("Sub-journey: Processing imported media", 3)
bullet("[P0] User can transcribe imported files using the same language and quality options as live recordings.")
bullet("[P0] User can view processing status (queued, processing, complete, failed).")
bullet("[P1] User can cancel an in-progress transcription job.")
bullet("[P2] User can auto-split long files into chapter segments.")

# Journey 3
h("User Journey 3: User organizing, searching, and reusing transcripts", 2)
p("Context: The transcript library is the retention engine. Users must be able to find past content quickly and export it into their existing workflows.")

h("Sub-journey: Library and search", 3)
bullet("[P0] User can view a list of all recordings/transcripts sorted by date.")
bullet("[P0] User can search across all transcripts by keyword and see matching results.")
bullet("[P0] User can open any saved transcript from the library.")
bullet("[P1] User can filter library items by date, language, or tag.")
bullet("[P2] User can archive or delete transcripts.")

h("Sub-journey: Export and sharing", 3)
bullet("[P0] User can export transcript as plain text (.txt).")
bullet("[P0] User can copy full transcript or selected text to clipboard.")
bullet("[P1] User can export transcript as DOCX or PDF.")
bullet("[P1] User can export subtitles as SRT or VTT with timestamps.")
bullet("[P2] User can share a read-only transcript link.")

# Journey 4
h("User Journey 4: User getting AI insights from long recordings", 2)
p("Context: Long meetings and lectures benefit from summaries and Q&A. This journey increases perceived value without requiring users to read full transcripts.")

h("Sub-journey: Summaries and highlights", 3)
bullet("[P1] User can generate an AI summary of a completed transcript.")
bullet("[P1] User can view bullet-point highlights of key moments.")
bullet("[P2] User can regenerate summary with different length (short vs detailed).")

h("Sub-journey: Chat with transcript", 3)
bullet("[P2] User can ask natural-language questions about a transcript (e.g., “What action items were mentioned?”).")
bullet("[P2] User can see answers grounded in transcript content with referenced timestamps.")

# Journey 5
h("User Journey 5: Privacy-conscious user choosing processing mode", 2)
p("Context: Different users have different privacy and connectivity needs. MVP should support at least one cloud path and lay groundwork for offline/local processing.")

h("Sub-journey: Processing preferences", 3)
bullet("[P0] User can transcribe via cloud API when online (default MVP path).")
bullet("[P1] User can see whether a transcript was processed locally or in the cloud.")
bullet("[P1] User can delete original audio after transcription while keeping text.")
bullet("[P2] User can transcribe fully on-device when offline/local model is available.")

# Journey 6
h("User Journey 6: Account, settings, and onboarding", 2)
p("Context: Users should reach first successful transcription quickly. Account creation can be deferred until sync or paid features are needed.")

h("Sub-journey: First-run experience", 3)
bullet("[P0] User can complete first transcription without signing up.")
bullet("[P0] User can grant microphone permission and see clear error states if denied.")
bullet("[P1] User can complete a short onboarding explaining record → transcribe → search → export.")
bullet("[P2] User can skip onboarding and access app immediately.")

h("Sub-journey: Account and sync", 3)
bullet("[P1] User can create an account to sync transcripts across devices.")
bullet("[P1] User can sign in with email or Google.")
bullet("[P2] User can manage subscription/billing for premium features.")

# 4. APPENDIX
h("4. APPENDIX", 1)
p("Competitive Reference Apps")
bullet("Whisper – Speech to Text (iOS): real-time transcription, 143+ languages, YouTube import, AI summaries, universal search.")
bullet("Whisper Transcription (iOS): share-extension import, local + cloud transcription, AI summaries, chat with transcript.")
bullet("Whisper Notes (iOS/macOS): offline-first, on-device Whisper Large V3 Turbo, subtitle export, privacy-focused.")

p("Technical Constraints (MVP)")
bullet("Transcription engine: OpenAI Whisper API for MVP cloud path; evaluate whisper.cpp or equivalent for future on-device support.")
bullet("Target platforms: Web app (primary) + desktop wrapper (Electron/Tauri) as stretch; mobile native deferred.")
bullet("Supported import formats MVP: MP3, M4A, WAV, MP4, MOV.")
bullet("Storage: Local IndexedDB/SQLite for transcript library; optional cloud sync in P1.")

p("Open Questions")
bullet("MVP pricing model: freemium with monthly transcription minutes vs one-time purchase vs fully free during beta?")
bullet("Should YouTube import be P1 or deferred due to legal/ToS complexity?")
bullet("Minimum language set for launch: English-only vs English + Spanish, French, German, Chinese?")
bullet("On-device transcription: ship in MVP as premium differentiator or defer to v2?")

p("Suggested Next Artifacts")
bullet("User flow wireframes: Record, Import, Library, Transcript Detail, Export.")
bullet("Technical architecture doc: audio pipeline, Whisper API integration, storage schema.")
bullet("Competitive feature matrix comparing Wisper MVP vs top 3 Whisper-class apps.")

doc.save(OUT)
print(f"Saved: {OUT}")
